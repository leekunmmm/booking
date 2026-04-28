# -*- coding: utf-8 -*-
"""Sinh tài liệu đặc tả yêu cầu & thiết kế (.docx) cho dự án LuxeHaven Booking.
   Phiên bản 2.0 – cập nhật đầy đủ theo codebase hiện tại (Voucher, MemberTier,
   PaymentMethod, Amenity, avatar, multi-image, admin đầy đủ).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(hdr[i], '1F497D')
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ─────────────────────────────────────────────────
# Tạo document
# ─────────────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===================================================
# TRANG BÌA
# ===================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('TRƯỜNG / KHOA CÔNG NGHỆ THÔNG TIN')
r.bold = True
r.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU & THIẾT KẾ')
r.bold = True
r.font.size = Pt(20)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('LuxeHaven – Hệ thống đặt phòng khách sạn')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

for _ in range(2):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Phiên bản: 2.0   |   Ngôn ngữ: Tiếng Việt   |   Tháng 4/2026')
r.italic = True

doc.add_page_break()

# ===================================================
# MỤC LỤC
# ===================================================
add_heading(doc, 'MỤC LỤC', level=1)
toc_items = [
    '1. Tổng quan dự án',
    '2. Yêu cầu chức năng (Functional Requirements)',
    '3. Yêu cầu phi chức năng (Non-Functional Requirements)',
    '4. Sơ đồ luồng dữ liệu (Data Flow Diagram – DFD)',
    '5. Sơ đồ Use Case',
    '6. Sơ đồ lớp (Class Diagram)',
    '7. Mô hình dữ liệu (Data Model / ERD)',
    '8. Thiết kế giao diện (Interface Design)',
    '9. Kế hoạch kiểm thử (Test Plan)',
]
for it in toc_items:
    doc.add_paragraph(it, style='List Number')
doc.add_page_break()

# ===================================================
# 1. TỔNG QUAN
# ===================================================
add_heading(doc, '1. Tổng quan dự án', level=1)
add_para(doc,
    'LuxeHaven là hệ thống web cho phép khách hàng tìm kiếm và đặt phòng khách sạn '
    'trực tuyến. Hệ thống cung cấp đầy đủ: xác thực người dùng, quản lý phòng (kèm '
    'hình ảnh và tiện nghi), đặt/huỷ/thanh toán phòng, áp dụng voucher giảm giá, '
    'phân hạng thành viên (Member Tier), và một bảng điều khiển Admin quản lý toàn bộ '
    'dữ liệu hệ thống.')

add_heading(doc, '1.1 Công nghệ sử dụng', level=2)
tech = [
    ('Backend', 'Java 17, Spring Boot 4.0.6, Spring JDBC (không dùng JPA/Hibernate)'),
    ('Frontend', 'Thymeleaf + Bootstrap 5 (CDN) + Bootstrap Icons'),
    ('CSDL', 'SQLite (file booking.db – không cần cài DB server)'),
    ('Build tool', 'Maven (wrapper mvnw / mvnw.cmd)'),
    ('Bảo mật', 'BCrypt (mật khẩu), HttpSession (phiên đăng nhập), AuthInterceptor'),
    ('Upload', 'Spring MultipartFile – lưu ảnh vào thư mục /uploads/'),
    ('Kiểm thử', 'JUnit 5 + Spring Boot Test (12 test tự động)'),
    ('IDE', 'VS Code / IntelliJ / Eclipse'),
    ('VCS', 'Git + GitHub'),
]
add_table(doc, ['Thành phần', 'Công nghệ / Thư viện'], tech, [4, 11])

add_heading(doc, '1.2 Phạm vi (Scope)', level=2)
add_para(doc, 'Trong phạm vi:', bold=True)
for s in [
    'Quản lý tài khoản: đăng ký, đăng nhập, đăng xuất, hồ sơ cá nhân, đổi avatar.',
    'Quản lý phòng: CRUD, upload nhiều ảnh, gán tiện nghi (Amenity), lọc nâng cao.',
    'Đặt phòng: chọn ngày, áp mã voucher, chọn phương thức thanh toán (MoMo / chuyển khoản).',
    'Thanh toán tượng trưng: xem QR, xác nhận → booking chuyển CONFIRMED.',
    'Huỷ booking, xem lịch sử booking cá nhân.',
    'Phân hạng thành viên tự động: BRONZE → SILVER → GOLD → DIAMOND → VIP.',
    'Admin: dashboard thống kê, quản lý phòng/booking/khách hàng/voucher, xác nhận thanh toán.',
]:
    doc.add_paragraph(s, style='List Bullet')

add_para(doc, 'Ngoài phạm vi:', bold=True)
for s in [
    'Cổng thanh toán thật (VNPay, Stripe, MoMo API).',
    'Email xác thực / khôi phục mật khẩu.',
    'Đa ngôn ngữ (chỉ tiếng Việt trong v2).',
    'Quản lý nhân viên, ca làm việc, kho hàng.',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ===================================================
# 2. YÊU CẦU CHỨC NĂNG
# ===================================================
add_heading(doc, '2. Yêu cầu chức năng (Functional Requirements)', level=1)

fr = [
    # Auth
    ('FR-01', 'Đăng ký tài khoản', 'Guest',
     'Nhập username, password (≥6 ký tự), họ tên, email, sđt → tạo tài khoản CUSTOMER. '
     'Mật khẩu hash BCrypt. Validate duplicate username/email.'),
    ('FR-02', 'Đăng nhập', 'Tất cả',
     'Xác thực username + password → lưu User vào HttpSession → redirect trang chủ.'),
    ('FR-03', 'Đăng xuất', 'Đã đăng nhập',
     'Huỷ session → redirect /login.'),
    ('FR-04', 'Hồ sơ cá nhân', 'Customer',
     'Xem và cập nhật họ tên / email / sđt. Upload ảnh đại diện (avatar). '
     'Hiển thị hạng thành viên và tổng chi tiêu.'),
    # Phòng
    ('FR-05', 'Xem danh sách phòng', 'Tất cả',
     'Hiển thị phòng AVAILABLE kèm ảnh, mô tả, giá, loại, tiện nghi.'),
    ('FR-06', 'Tìm phòng theo ngày', 'Tất cả',
     'Lọc loại bỏ phòng đã có booking PENDING/CONFIRMED giao với [check_in, check_out].'),
    ('FR-07', 'Lọc phòng theo loại', 'Tất cả',
     'Lọc theo STANDARD / DELUXE / SUITE.'),
    ('FR-08', 'Xem chi tiết phòng', 'Tất cả',
     'Gallery nhiều ảnh, mô tả, sức chứa, danh sách tiện nghi (icon + tên), giá/đêm.'),
    # Booking
    ('FR-09', 'Đặt phòng', 'Customer',
     'Nhập check_in, check_out, thông tin liên hệ, ghi chú. Kiểm tra xung đột lịch. '
     'Tính total_price = price_per_night × số_đêm, trừ discount voucher. Tạo booking PENDING.'),
    ('FR-10', 'Áp dụng voucher', 'Customer',
     'Nhập mã voucher khi đặt phòng. Hệ thống kiểm tra: active, còn hạn, chưa hết lượt, '
     'đúng tier thành viên → giảm % tổng tiền.'),
    ('FR-11', 'Thanh toán tượng trưng', 'Customer (chủ booking)',
     'Chọn phương thức (MoMo / chuyển khoản), xem QR tương ứng → xác nhận → '
     'payment_status = PAID, status = CONFIRMED.'),
    ('FR-12', 'Xem lịch sử đặt phòng', 'Customer',
     'Bảng danh sách booking của user, mới nhất trước, kèm trạng thái và tổng tiền.'),
    ('FR-13', 'Huỷ đặt phòng', 'Customer (chủ booking)',
     'Cho huỷ khi status = PENDING hoặc CONFIRMED, check_in chưa diễn ra → status = CANCELLED.'),
    # Admin – phòng
    ('FR-14', 'Quản lý phòng (CRUD)', 'Admin',
     'Tạo/sửa/xoá phòng. Upload nhiều ảnh (lưu thư mục /uploads/). '
     'Gán tiện nghi từ danh sách cố định (Amenity enum). Không xoá phòng có booking đang mở.'),
    # Admin – booking
    ('FR-15', 'Xem tất cả booking', 'Admin',
     'Bảng booking toàn hệ thống, lọc theo status / payment_status / tên khách.'),
    ('FR-16', 'Cập nhật trạng thái booking', 'Admin',
     'Đổi status: PENDING→CONFIRMED→COMPLETED hoặc CANCELLED.'),
    ('FR-17', 'Xác nhận thanh toán', 'Admin',
     'Admin bấm "Xác nhận thanh toán" → payment_status = PAID, status = CONFIRMED.'),
    # Admin – voucher
    ('FR-18', 'Quản lý voucher (CRUD)', 'Admin',
     'Tạo/sửa/xoá mã voucher: code, discount_percent, valid_until, max_uses, '
     'target_tier (null = tất cả, hoặc chỉ BRONZE/SILVER/GOLD/DIAMOND/VIP).'),
    # Admin – khách hàng
    ('FR-19', 'Quản lý khách hàng', 'Admin',
     'Xem danh sách khách hàng kèm hạng thành viên, tổng chi tiêu. '
     'Lọc theo tên/email hoặc tier. Xoá tài khoản khách hàng.'),
    # Member tier
    ('FR-20', 'Phân hạng thành viên', 'Hệ thống tự động',
     'Tính tier dựa trên tổng chi tiêu (booking COMPLETED + CONFIRMED + PAID): '
     '< 5 tr = BRONZE, 5–15 tr = SILVER, 15–30 tr = GOLD, 30–60 tr = DIAMOND, ≥ 60 tr = VIP.'),
]
add_table(doc, ['ID', 'Tên chức năng', 'Vai trò', 'Mô tả'], fr, [1.5, 3.5, 3, 8])

doc.add_page_break()

# ===================================================
# 3. YÊU CẦU PHI CHỨC NĂNG
# ===================================================
add_heading(doc, '3. Yêu cầu phi chức năng (Non-Functional Requirements)', level=1)
nfr = [
    ('NFR-01', 'Hiệu năng',   'Trang danh sách phòng tải < 2 giây với 100 phòng.'),
    ('NFR-02', 'Bảo mật',     'Mật khẩu hash BCrypt (cost 10). Không lưu plaintext.'),
    ('NFR-03', 'Bảo mật',     'AuthInterceptor phân quyền: CUSTOMER bị chặn /admin/**; chưa đăng nhập bị redirect /login.'),
    ('NFR-04', 'Bảo mật',     'Kiểm tra ownership: customer chỉ thanh toán/huỷ booking của chính mình.'),
    ('NFR-05', 'Bảo mật',     'Session timeout 30 phút không hoạt động (server.servlet.session.timeout=1800).'),
    ('NFR-06', 'Khả dụng',    'Giao diện responsive Bootstrap 5, dùng được trên mobile/tablet/desktop.'),
    ('NFR-07', 'Tin cậy',     'Validate đầu vào server-side: check_out > check_in, ngày không quá khứ, không trùng booking.'),
    ('NFR-08', 'Tin cậy',     'BusinessException tập trung, thông báo lỗi thân thiện bằng tiếng Việt.'),
    ('NFR-09', 'Bảo trì',     'Phân lớp rõ: model / repository / service / controller / config. Unit test coverage ≥ 60%.'),
    ('NFR-10', 'Lưu trữ',     'Ảnh upload lưu vào thư mục /uploads/ local, phục vụ qua /uploads/** static resource.'),
    ('NFR-11', 'Tương thích', 'Chạy được trên Windows / Linux / macOS với Java 17+.'),
    ('NFR-12', 'Triển khai',  'Build và chạy chỉ cần: ./mvnw spring-boot:run – không cần cài thêm DB server.'),
]
add_table(doc, ['ID', 'Loại', 'Mô tả'], nfr, [1.5, 2.5, 12])

doc.add_page_break()

# ===================================================
# 4. DFD
# ===================================================
add_heading(doc, '4. Sơ đồ luồng dữ liệu (Data Flow Diagram – DFD)', level=1)
add_para(doc,
    'Lưu ý: Các sơ đồ ASCII dưới đây mô tả cấu trúc luồng dữ liệu. '
    'Phiên bản đồ hoạ chính thức được vẽ trên draw.io (xem file diagrams/dfd.drawio).',
    italic=True)

add_heading(doc, '4.1 DFD mức 0 – Context Diagram', level=2)
add_para(doc,
    'Ở mức 0, hệ thống là một hộp đen nhận yêu cầu từ hai tác nhân ngoài: '
    'Khách hàng (Guest/Customer) và Quản trị viên (Admin).')

add_code_block(doc, """\
   ┌─────────────┐     yêu cầu đặt phòng / xem phòng      ┌───────────────────────────┐
   │             │ ─────────────────────────────────────→  │                           │
   │   Khách     │                                         │   Hệ thống LuxeHaven      │
   │   hàng      │ ←─────────────────────────────────────  │   Booking                 │
   │             │     xác nhận đặt / hoá đơn / thông báo  │                           │
   └─────────────┘                                         │                           │
                                                           │                           │
   ┌─────────────┐   thao tác quản lý (CRUD phòng,         │                           │
   │  Quản trị   │   booking, voucher, khách hàng)          │                           │
   │  viên       │ ─────────────────────────────────────→  │                           │
   │  (Admin)    │ ←─────────────────────────────────────  │                           │
   └─────────────┘     báo cáo / danh sách / thống kê       └───────────────────────────┘
""", caption='Hình 4.1 – DFD mức 0 (Context Diagram)')

add_heading(doc, '4.2 DFD mức 1', level=2)
add_para(doc,
    'Phân rã thành 5 quy trình chính, 4 kho dữ liệu: users, rooms, bookings, vouchers.')

add_code_block(doc, """\
                          ┌────────────────────┐
  Khách ─đăng ký/đn──→   │ 1. Quản lý         │ ──→ [D1: users]
                          │    tài khoản       │
                          └────────────────────┘
                                    │
                                    ↓
                          ┌────────────────────┐
  Khách ─tìm/xem phòng→  │ 2. Tìm kiếm &      │ ←── [D2: rooms]
                          │    xem phòng       │
                          └────────────────────┘
                                    │
                                    ↓
                          ┌────────────────────┐
  Khách ─đặt/thanh toán→ │ 3. Đặt phòng &     │ ──→ [D3: bookings]
                          │    thanh toán      │ ←── [D4: vouchers]
                          └────────────────────┘
                                    │
                                    ↓
                          ┌────────────────────┐
  Khách ─huỷ/xem ls──→   │ 4. Quản lý         │ ←── [D3: bookings]
                          │    booking cá nhân │
                          └────────────────────┘
                                    │
  Admin ─CRUD / xem──→   ┌────────────────────┐
                          │ 5. Quản trị        │ ←── [D1][D2][D3][D4]
                          │    hệ thống        │
                          └────────────────────┘
""", caption='Hình 4.2 – DFD mức 1')

add_heading(doc, '4.3 Mô tả kho dữ liệu', level=2)
stores = [
    ('D1: users',    'Thông tin tài khoản: username, password_hash, role, avatar, tier tính toán.'),
    ('D2: rooms',    'Thông tin phòng: số phòng, loại, giá, mô tả, danh sách ảnh, tiện nghi.'),
    ('D3: bookings', 'Đơn đặt phòng: user_id, room_id, ngày, tổng tiền, trạng thái, phương thức thanh toán.'),
    ('D4: vouchers', 'Mã giảm giá: code, %, hạn dùng, số lượt, target_tier.'),
]
add_table(doc, ['Kho dữ liệu', 'Nội dung'], stores, [3, 13])

doc.add_page_break()

# ===================================================
# 5. USE CASE
# ===================================================
add_heading(doc, '5. Sơ đồ Use Case', level=1)

add_heading(doc, '5.1 Các tác nhân (Actors)', level=2)
actors = [
    ('Guest',     'Người dùng chưa đăng nhập. Có thể xem danh sách phòng, tìm kiếm, chi tiết phòng, đăng ký tài khoản.'),
    ('Customer',  'Đã đăng nhập với role CUSTOMER. Bao gồm mọi quyền Guest + đặt phòng, huỷ, thanh toán, xem lịch sử, hồ sơ cá nhân.'),
    ('Admin',     'Role ADMIN. Toàn quyền: quản lý phòng, booking, voucher, khách hàng, xem dashboard thống kê.'),
    ('Hệ thống',  'Tính tier thành viên tự động dựa trên tổng chi tiêu, cập nhật usedCount voucher.'),
]
add_table(doc, ['Actor', 'Mô tả'], actors, [2.5, 13])

add_heading(doc, '5.2 Danh sách Use Case', level=2)
ucs = [
    ('UC-01', 'Đăng ký tài khoản',          'Guest',                    'FR-01'),
    ('UC-02', 'Đăng nhập',                   'Guest / Customer / Admin',  'FR-02'),
    ('UC-03', 'Đăng xuất',                   'Customer / Admin',          'FR-03'),
    ('UC-04', 'Xem & cập nhật hồ sơ',        'Customer',                  'FR-04'),
    ('UC-05', 'Xem danh sách phòng',          'Tất cả',                   'FR-05'),
    ('UC-06', 'Tìm phòng theo ngày/loại',     'Tất cả',                   'FR-06, FR-07'),
    ('UC-07', 'Xem chi tiết phòng',           'Tất cả',                   'FR-08'),
    ('UC-08', 'Đặt phòng',                    'Customer',                  'FR-09'),
    ('UC-09', 'Áp dụng voucher',              'Customer',                  'FR-10'),
    ('UC-10', 'Thanh toán tượng trưng',       'Customer',                  'FR-11'),
    ('UC-11', 'Xem lịch sử đặt phòng',       'Customer',                  'FR-12'),
    ('UC-12', 'Huỷ đặt phòng',               'Customer',                  'FR-13'),
    ('UC-13', 'Quản lý phòng (CRUD)',         'Admin',                     'FR-14'),
    ('UC-14', 'Xem tất cả booking',           'Admin',                     'FR-15'),
    ('UC-15', 'Cập nhật trạng thái booking',  'Admin',                     'FR-16'),
    ('UC-16', 'Xác nhận thanh toán',          'Admin',                     'FR-17'),
    ('UC-17', 'Quản lý voucher (CRUD)',        'Admin',                     'FR-18'),
    ('UC-18', 'Quản lý khách hàng',           'Admin',                     'FR-19'),
    ('UC-19', 'Tính tier thành viên',         'Hệ thống',                  'FR-20'),
]
add_table(doc, ['ID', 'Use Case', 'Actor', 'FR liên quan'], ucs, [1.5, 5.5, 4, 5])

add_heading(doc, '5.3 Sơ đồ Use Case (mô tả văn bản)', level=2)
add_para(doc,
    'Ghi chú: Phiên bản đồ hoạ chính thức xem file diagrams/usecase.drawio.',
    italic=True)
add_code_block(doc, """\
  ┌─────────────────────────────────────────────────────────────────────┐
  │                        Hệ thống LuxeHaven                           │
  │                                                                     │
  │  ( UC-01 Đăng ký )           ( UC-05 Xem phòng )                   │
  │  ( UC-02 Đăng nhập )         ( UC-06 Tìm phòng theo ngày/loại )    │
  │  ( UC-03 Đăng xuất )         ( UC-07 Chi tiết phòng )              │
  │  ( UC-04 Hồ sơ cá nhân )                                           │
  │                                                                     │
  │  ( UC-08 Đặt phòng )         ( UC-13 CRUD phòng )                  │
  │  ( UC-09 Áp voucher )        ( UC-14 Xem tất cả booking )          │
  │  ( UC-10 Thanh toán )        ( UC-15 Cập nhật trạng thái )         │
  │  ( UC-11 Lịch sử )           ( UC-16 Xác nhận thanh toán )         │
  │  ( UC-12 Huỷ phòng )         ( UC-17 Quản lý voucher )             │
  │                               ( UC-18 Quản lý khách hàng )         │
  │  ( UC-19 Tính tier – Hệ thống tự động )                            │
  └───────────▲──────────────────────────────▲────────────────────────┘
              │                              │
  ┌───────────┴──────────┐        ┌──────────┴──────────┐
  │        Guest         │        │        Admin         │
  └───────────▲──────────┘        └─────────────────────┘
              │ <<extends>>
  ┌───────────┴──────────┐
  │      Customer        │
  └─────────────────────-┘
""", caption='Hình 5.1 – Sơ đồ Use Case tổng quát')

add_heading(doc, '5.4 Đặc tả chi tiết UC-08 (Đặt phòng)', level=2)
uc_detail = [
    ('Tên',             'UC-08 Đặt phòng'),
    ('Actor chính',     'Customer'),
    ('Tiền điều kiện',  'User đã đăng nhập (role CUSTOMER). Phòng tồn tại và không có booking trùng ngày.'),
    ('Hậu điều kiện',   'Booking mới tạo với status = PENDING, usedCount voucher +1 (nếu có).'),
    ('Luồng chính',
        '1. Customer chọn phòng → bấm "Đặt ngay".\n'
        '2. Hệ thống hiện form: check_in, check_out, họ tên liên hệ, sđt, email, ghi chú, mã voucher.\n'
        '3. Customer điền và submit.\n'
        '4. Hệ thống validate ngày (không quá khứ, check_out > check_in).\n'
        '5. Kiểm tra xung đột booking.\n'
        '6. Nếu có voucher: xác thực và tính discount.\n'
        '7. Tính total_price = nights × price_per_night × (1 - discount%).\n'
        '8. Tạo booking PENDING, redirect /my-bookings.'),
    ('Alt 4a', 'Ngày không hợp lệ → báo lỗi → quay bước 3.'),
    ('Alt 5a', 'Phòng đã có booking trùng → báo lỗi → quay bước 3.'),
    ('Alt 6a', 'Voucher không hợp lệ / hết lượt / sai tier → báo lỗi, tiếp tục không voucher.'),
]
add_table(doc, ['Trường', 'Nội dung'], uc_detail, [3.5, 12.5])

doc.add_page_break()

# ===================================================
# 6. CLASS DIAGRAM
# ===================================================
add_heading(doc, '6. Sơ đồ lớp (Class Diagram)', level=1)
add_para(doc,
    'Hệ thống có 4 entity chính (User, Room, Booking, Voucher) và 6 enum hỗ trợ. '
    'Phiên bản đồ hoạ chính thức xem file diagrams/classdiagram.drawio.',
    italic=True)

add_code_block(doc, """\
┌──────────────────────────┐        ┌───────────────────────────────┐
│           User            │        │              Room              │
├──────────────────────────┤        ├───────────────────────────────┤
│ - id : Long               │        │ - id : Long                   │
│ - username : String       │        │ - roomNumber : String         │
│ - password : String       │        │ - type : RoomType             │
│ - fullName : String       │        │ - pricePerNight : double      │
│ - email : String          │        │ - capacity : int              │
│ - phone : String          │        │ - description : String        │
│ - role : Role             │        │ - status : RoomStatus         │
│ - avatarUrl : String      │        │ - imagesRaw : String          │
│ - memberTier : String     │        │ - amenitiesRaw : String       │
│ - createdAt : LocalDT     │        ├───────────────────────────────┤
├──────────────────────────┤        │ + getImages() : List<String>  │
│ + isAdmin() : boolean     │        │ + getAmenities() : List<Amn>  │
└───────────┬──────────────┘        │ + getFirstImage() : String    │
            │ 1                     └───────────────┬───────────────┘
            │                                       │ 1
            │ *   ┌──────────────────────────────── ┘
            └────→│           Booking               │
                  ├─────────────────────────────────┤
                  │ - id : Long                     │
                  │ - userId : Long                 │
                  │ - roomId : Long                 │
                  │ - checkIn : LocalDate           │
                  │ - checkOut : LocalDate          │
                  │ - totalPrice : double           │
                  │ - status : BookingStatus        │
                  │ - paymentStatus : PaymentStatus │
                  │ - paymentMethod : PaymentMethod │
                  │ - customerFullName : String     │
                  │ - customerPhone : String        │
                  │ - customerEmail : String        │
                  │ - customerNote : String         │
                  │ - voucherCode : String          │
                  │ - createdAt : LocalDateTime     │
                  ├─────────────────────────────────┤
                  │ + getNumberOfNights() : long    │
                  └─────────────────────────────────┘

┌──────────────────────────┐
│          Voucher          │
├──────────────────────────┤
│ - id : Long               │
│ - code : String           │
│ - discountPercent : int   │
│ - validUntil : LocalDate  │
│ - maxUses : int           │
│ - usedCount : int         │
│ - active : boolean        │
│ - targetTier : String     │
├──────────────────────────┤
│ + isUsable() : boolean    │
└──────────────────────────┘

Enums:
  Role           = { CUSTOMER, ADMIN }
  RoomType       = { STANDARD, DELUXE, SUITE }
  RoomStatus     = { AVAILABLE, MAINTENANCE }
  BookingStatus  = { PENDING, CONFIRMED, CANCELLED, COMPLETED }
  PaymentStatus  = { UNPAID, PAID }
  PaymentMethod  = { MOMO("Ví MoMo"), BANK_TRANSFER("Chuyển khoản ngân hàng") }
  MemberTier     = { BRONZE, SILVER, GOLD, DIAMOND, VIP }
  Amenity        = { WIFI, AC, TV, MINIBAR, JACUZZI, BREAKFAST, BUTLER, POOL, GYM, SPA }
""", caption='Hình 6.1 – Class Diagram')

add_heading(doc, '6.1 Quan hệ giữa các lớp', level=2)
rel = [
    ('User – Booking',    '1 .. *', 'Một user có nhiều booking; mỗi booking thuộc đúng 1 user (userId FK).'),
    ('Room – Booking',    '1 .. *', 'Một phòng được đặt nhiều lần (khác thời gian); mỗi booking thuộc 1 phòng (roomId FK).'),
    ('Booking – Voucher', '* .. 0..1', 'Mỗi booking tuỳ chọn dùng 1 voucher (lưu voucherCode). Voucher không FK cứng.'),
]
add_table(doc, ['Quan hệ', 'Bội số', 'Ý nghĩa'], rel, [3.5, 2, 10.5])

doc.add_page_break()

# ===================================================
# 7. DATA MODEL / ERD
# ===================================================
add_heading(doc, '7. Mô hình dữ liệu (Data Model / ERD)', level=1)
add_para(doc,
    'Sử dụng SQLite. Có 4 bảng: users, rooms, bookings, vouchers. '
    'Phiên bản sơ đồ ERD đồ hoạ xem file diagrams/erd.drawio.',
    italic=True)

add_heading(doc, '7.1 Bảng users', level=2)
users_cols = [
    ('id',         'INTEGER', 'PK, AUTOINCREMENT',              'Khoá chính'),
    ('username',   'TEXT',    'NOT NULL, UNIQUE',               'Tên đăng nhập (duy nhất)'),
    ('password',   'TEXT',    'NOT NULL',                       'Mật khẩu hash BCrypt'),
    ('full_name',  'TEXT',    'NOT NULL',                       'Họ và tên đầy đủ'),
    ('email',      'TEXT',    'NOT NULL, UNIQUE',               'Email (duy nhất)'),
    ('phone',      'TEXT',    '',                               'Số điện thoại'),
    ('role',       'TEXT',    "NOT NULL DEFAULT 'CUSTOMER'",    'CUSTOMER | ADMIN'),
    ('avatar_url', 'TEXT',    '',                               'Đường dẫn ảnh đại diện (null = default)'),
    ('created_at', 'TEXT',    'NOT NULL DEFAULT CURRENT_TIMESTAMP', 'Thời điểm tạo tài khoản'),
]
add_table(doc, ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'], users_cols, [3, 2, 5, 6])

add_heading(doc, '7.2 Bảng rooms', level=2)
rooms_cols = [
    ('id',               'INTEGER', 'PK, AUTOINCREMENT',              'Khoá chính'),
    ('room_number',      'TEXT',    'NOT NULL, UNIQUE',               'Số phòng (vd: 101)'),
    ('type',             'TEXT',    'NOT NULL',                       'STANDARD | DELUXE | SUITE'),
    ('price_per_night',  'REAL',    'NOT NULL, CHECK > 0',           'Giá/đêm (VND)'),
    ('capacity',         'INTEGER', 'NOT NULL, CHECK > 0',           'Số người tối đa'),
    ('description',      'TEXT',    '',                               'Mô tả phòng'),
    ('status',           'TEXT',    "NOT NULL DEFAULT 'AVAILABLE'",   'AVAILABLE | MAINTENANCE'),
    ('images_raw',       'TEXT',    '',                               'URL ảnh, phân cách bằng dấu phẩy'),
    ('amenities_raw',    'TEXT',    '',                               'Tên Amenity enum, phân cách bằng dấu phẩy'),
]
add_table(doc, ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'], rooms_cols, [3, 2, 5, 6])

add_heading(doc, '7.3 Bảng bookings', level=2)
bk_cols = [
    ('id',                 'INTEGER', 'PK, AUTOINCREMENT',              'Khoá chính'),
    ('user_id',            'INTEGER', 'NOT NULL, FK → users(id)',        'Người đặt phòng'),
    ('room_id',            'INTEGER', 'NOT NULL, FK → rooms(id)',        'Phòng được đặt'),
    ('check_in',           'DATE',    'NOT NULL',                        'Ngày nhận phòng'),
    ('check_out',          'DATE',    'NOT NULL, CHECK > check_in',      'Ngày trả phòng'),
    ('total_price',        'REAL',    'NOT NULL',                        'Tổng tiền (đã trừ voucher)'),
    ('status',             'TEXT',    "NOT NULL DEFAULT 'PENDING'",      'PENDING|CONFIRMED|CANCELLED|COMPLETED'),
    ('payment_status',     'TEXT',    "NOT NULL DEFAULT 'UNPAID'",       'UNPAID | PAID'),
    ('payment_method',     'TEXT',    '',                                'MOMO | BANK_TRANSFER'),
    ('customer_full_name', 'TEXT',    '',                                'Tên liên hệ khi đặt'),
    ('customer_phone',     'TEXT',    '',                                'SĐT liên hệ khi đặt'),
    ('customer_email',     'TEXT',    '',                                'Email liên hệ khi đặt'),
    ('customer_note',      'TEXT',    '',                                'Ghi chú thêm'),
    ('voucher_code',       'TEXT',    '',                                'Mã voucher đã áp dụng (null nếu không có)'),
    ('created_at',         'TEXT',    'NOT NULL DEFAULT CURRENT_TIMESTAMP', 'Thời điểm đặt phòng'),
]
add_table(doc, ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'], bk_cols, [3, 2, 5, 6])

add_heading(doc, '7.4 Bảng vouchers', level=2)
vc_cols = [
    ('id',               'INTEGER', 'PK, AUTOINCREMENT',    'Khoá chính'),
    ('code',             'TEXT',    'NOT NULL, UNIQUE',     'Mã voucher (in hoa)'),
    ('discount_percent', 'INTEGER', 'NOT NULL, CHECK 1–99', 'Phần trăm giảm giá'),
    ('valid_until',      'DATE',    '',                     'Hạn sử dụng (null = không hạn)'),
    ('max_uses',         'INTEGER', 'NOT NULL, DEFAULT 1',  'Số lượt dùng tối đa'),
    ('used_count',       'INTEGER', 'NOT NULL, DEFAULT 0',  'Số lượt đã dùng'),
    ('active',           'INTEGER', 'NOT NULL, DEFAULT 1',  '1=bật, 0=tắt'),
    ('target_tier',      'TEXT',    '',                     'null=tất cả | BRONZE|SILVER|GOLD|DIAMOND|VIP'),
]
add_table(doc, ['Cột', 'Kiểu', 'Ràng buộc', 'Mô tả'], vc_cols, [3, 2, 4, 7])

add_heading(doc, '7.5 Sơ đồ quan hệ (ERD – mô tả văn bản)', level=2)
add_code_block(doc, """\
  ┌──────────┐       ┌────────────────┐       ┌──────────┐
  │  users   │ 1── * │    bookings    │ *──1  │  rooms   │
  │──────────│       │────────────────│       │──────────│
  │ id  (PK) │       │ id  (PK)       │       │ id  (PK) │
  │ username │       │ user_id  (FK)──┘       │ room_num │
  │ password │       │ room_id  (FK)──────────│ type     │
  │ full_name│       │ check_in               │ price    │
  │ email    │       │ check_out              │ capacity │
  │ phone    │       │ total_price            │ status   │
  │ role     │       │ status                 │ images   │
  │ avatar   │       │ payment_status         │ amenity  │
  └──────────┘       │ payment_method         └──────────┘
                     │ customer_*
                     │ voucher_code ──────────→ ┌──────────────┐
                     │ created_at              │   vouchers   │
                     └─────────────────────────│──────────────│
                                               │ id  (PK)     │
                                               │ code (UNIQUE)│
                                               │ discount_%   │
                                               │ valid_until  │
                                               │ max_uses     │
                                               │ used_count   │
                                               │ active       │
                                               │ target_tier  │
                                               └──────────────┘

  Ghi chú: voucher_code là text reference (không phải FK cứng) để tránh ràng buộc
  khi admin xoá voucher sau khi đã được sử dụng.
""", caption='Hình 7.1 – ERD (Entity Relationship Diagram)')

add_heading(doc, '7.6 Câu lệnh DDL đầy đủ', level=2)
add_code_block(doc, """\
CREATE TABLE IF NOT EXISTS users (
  id          INTEGER PRIMARY KEY AUTOINCREMENT,
  username    TEXT    NOT NULL UNIQUE,
  password    TEXT    NOT NULL,
  full_name   TEXT    NOT NULL,
  email       TEXT    NOT NULL UNIQUE,
  phone       TEXT,
  role        TEXT    NOT NULL DEFAULT 'CUSTOMER',
  avatar_url  TEXT,
  created_at  TEXT    NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE IF NOT EXISTS rooms (
  id              INTEGER PRIMARY KEY AUTOINCREMENT,
  room_number     TEXT    NOT NULL UNIQUE,
  type            TEXT    NOT NULL,
  price_per_night REAL    NOT NULL CHECK (price_per_night > 0),
  capacity        INTEGER NOT NULL CHECK (capacity > 0),
  description     TEXT,
  status          TEXT    NOT NULL DEFAULT 'AVAILABLE',
  images_raw      TEXT,
  amenities_raw   TEXT
);

CREATE TABLE IF NOT EXISTS bookings (
  id                 INTEGER PRIMARY KEY AUTOINCREMENT,
  user_id            INTEGER NOT NULL,
  room_id            INTEGER NOT NULL,
  check_in           DATE    NOT NULL,
  check_out          DATE    NOT NULL,
  total_price        REAL    NOT NULL,
  status             TEXT    NOT NULL DEFAULT 'PENDING',
  payment_status     TEXT    NOT NULL DEFAULT 'UNPAID',
  payment_method     TEXT,
  customer_full_name TEXT,
  customer_phone     TEXT,
  customer_email     TEXT,
  customer_note      TEXT,
  voucher_code       TEXT,
  created_at         TEXT    NOT NULL DEFAULT CURRENT_TIMESTAMP,
  FOREIGN KEY (user_id) REFERENCES users(id),
  FOREIGN KEY (room_id) REFERENCES rooms(id),
  CHECK (check_out > check_in)
);

CREATE TABLE IF NOT EXISTS vouchers (
  id               INTEGER PRIMARY KEY AUTOINCREMENT,
  code             TEXT    NOT NULL UNIQUE,
  discount_percent INTEGER NOT NULL CHECK (discount_percent BETWEEN 1 AND 99),
  valid_until      DATE,
  max_uses         INTEGER NOT NULL DEFAULT 1,
  used_count       INTEGER NOT NULL DEFAULT 0,
  active           INTEGER NOT NULL DEFAULT 1,
  target_tier      TEXT
);
""")

doc.add_page_break()

# ===================================================
# 8. INTERFACE DESIGN
# ===================================================
add_heading(doc, '8. Thiết kế giao diện (Interface Design)', level=1)
add_para(doc,
    'Giao diện Web được render bằng Thymeleaf, dùng Bootstrap 5 + Bootstrap Icons. '
    'Layout chia sẻ qua fragments/layout.html (navbar + footer). '
    'Màu chủ đạo: xanh navy #1F497D (header/btn primary). Responsive mobile-first.')

add_heading(doc, '8.1 Danh sách trang', level=2)
pages = [
    ('/', 'Trang chủ (home.html)',
     'Hero banner, form tìm phòng nhanh (check-in, check-out, loại), grid phòng nổi bật, section dịch vụ.'),
    ('/rooms', 'Danh sách phòng (rooms.html)',
     'Card grid 3 cột responsive. Mỗi card: ảnh đầu, tên phòng, loại badge, tiện nghi icons, giá, nút "Xem chi tiết".'),
    ('/rooms/{id}', 'Chi tiết phòng (room-detail.html)',
     'Gallery nhiều ảnh (carousel), mô tả đầy đủ, sức chứa, danh sách tiện nghi có icon, giá/đêm, nút "Đặt ngay".'),
    ('/bookings/new', 'Form đặt phòng (booking-form.html)',
     'Tóm tắt phòng, chọn check-in/out, họ tên/sđt/email liên hệ, ghi chú, ô nhập mã voucher, hiển thị tổng tiền động.'),
    ('/bookings/{id}/payment', 'Trang thanh toán (booking-payment.html)',
     'Tóm tắt booking, chọn phương thức (MoMo / chuyển khoản ngân hàng), hiển thị QR tương ứng, nút "Xác nhận đã chuyển".'),
    ('/my-bookings', 'Lịch sử đặt phòng (my-bookings.html)',
     'Bảng booking của user: mã, phòng, ngày, tổng tiền, trạng thái badge, nút "Thanh toán" / "Huỷ".'),
    ('/profile', 'Hồ sơ cá nhân (profile.html)',
     'Ảnh avatar (có nút đổi ảnh), thông tin cá nhân chỉnh sửa inline, hạng thành viên + tổng chi tiêu.'),
    ('/login', 'Đăng nhập (login.html)',
     'Form username + password, link sang trang đăng ký.'),
    ('/register', 'Đăng ký (register.html)',
     'Form: username, password, fullName, email, phone. Validate cả client và server.'),
    ('/contact', 'Liên hệ (contact.html)',
     'Thông tin liên hệ khách sạn (địa chỉ, sđt, email, bản đồ nhúng).'),
    ('/services', 'Dịch vụ (services.html)',
     'Giới thiệu các dịch vụ của khách sạn (spa, nhà hàng, hồ bơi...).'),
    ('/admin', 'Dashboard Admin (admin/dashboard.html)',
     'Thống kê: tổng phòng, tổng khách, booking PENDING/CONFIRMED/COMPLETED, tổng doanh thu, bảng booking gần nhất.'),
    ('/admin/rooms', 'Quản lý phòng (admin/rooms.html)',
     'Bảng phòng có filter (loại, trạng thái, tìm kiếm). Nút Thêm/Sửa/Xoá.'),
    ('/admin/rooms/new|/{id}/edit', 'Form phòng (admin/room-form.html)',
     'Fields: số phòng, loại, giá, sức chứa, mô tả, trạng thái. Upload ảnh AJAX + preview. Checkbox tiện nghi.'),
    ('/admin/bookings', 'Quản lý booking (admin/bookings.html)',
     'Bảng booking toàn hệ thống. Filter: status, payment_status, tìm theo tên. Nút cập nhật trạng thái / xác nhận thanh toán.'),
    ('/admin/vouchers', 'Quản lý voucher (admin/vouchers.html)',
     'Bảng voucher. Filter: code, active, tier. Nút Thêm/Sửa/Xoá.'),
    ('/admin/vouchers/new|/{id}/edit', 'Form voucher (admin/voucher-form.html)',
     'Fields: code, discount%, hạn dùng, số lượt, active toggle, target tier.'),
    ('/admin/customers', 'Quản lý khách hàng (admin/customers.html)',
     'Bảng user role CUSTOMER. Hiển thị hạng thành viên, tổng chi tiêu. Filter: tên/email, tier. Nút Xoá.'),
]
add_table(doc, ['URL', 'Trang', 'Mô tả layout'], pages, [4.5, 4, 7.5])

add_heading(doc, '8.2 Quy ước giao diện', level=2)
for s in [
    'Màu chủ đạo: xanh navy #1F497D – dùng cho navbar, nút primary, badge admin.',
    'Font: Bootstrap default (Segoe UI / Calibri trên Windows).',
    'Responsive: navbar collapse ở < 992px, card/bảng scroll ngang trên mobile.',
    'Thông báo: alert-success (xanh lá) cho thành công, alert-danger (đỏ) cho lỗi.',
    'Format ngày: dd/MM/yyyy. Format tiền: 1.234.567 VND (NumberFormat tiếng Việt).',
    'Ảnh phòng: mặc định dùng ảnh Unsplash nếu chưa upload.',
    'Bootstrap Icons (bi-*) dùng cho tiện nghi phòng và icon nav.',
    'Flash messages (RedirectAttributes) tự tắt sau khi render.',
]:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '8.3 Phân hạng thành viên – bảng màu', level=2)
tiers = [
    ('BRONZE',  '< 5.000.000 VND',    '🥉 Bronze',  'text-warning (nâu đồng)'),
    ('SILVER',  '5 – 15 triệu VND',   '🥈 Silver',  'text-secondary (bạc)'),
    ('GOLD',    '15 – 30 triệu VND',  '🥇 Gold',    'text-warning (vàng)'),
    ('DIAMOND', '30 – 60 triệu VND',  '💎 Diamond', 'text-info (xanh)'),
    ('VIP',     '≥ 60 triệu VND',     '👑 VIP',     'text-danger (đỏ)'),
]
add_table(doc, ['Hạng', 'Tổng chi tiêu', 'Hiển thị', 'Màu Bootstrap'], tiers, [2.5, 4, 3, 6.5])

doc.add_page_break()

# ===================================================
# 9. TEST PLAN
# ===================================================
add_heading(doc, '9. Kế hoạch kiểm thử (Test Plan)', level=1)
add_para(doc,
    'Bộ kiểm thử gồm 2 phần: (1) Unit test tự động với JUnit 5; '
    '(2) Test case thủ công theo bảng dưới. '
    'File test-cases.md trong thư mục docs/ chứa đầy đủ bộ 30 test case để copy vào Excel.')

add_heading(doc, '9.1 Unit test tự động', level=2)
ut = [
    ('BookingApplicationTests',        '1', 'Kiểm tra Spring context khởi động thành công.'),
    ('service/AuthServiceTest',        '5', 'Đăng ký (thành công, username trùng, mật khẩu ngắn, email sai), hash BCrypt.'),
    ('service/BookingServiceTest',     '6', 'Tính tổng tiền, kiểm tra trùng lịch, thanh toán, huỷ, ownership check.'),
]
add_table(doc, ['File Test', 'Số TC', 'Mô tả'], ut, [6, 1.5, 8.5])
add_para(doc, 'Chạy lệnh: ./mvnw test', italic=True)

add_heading(doc, '9.2 Test case thủ công (tóm tắt)', level=2)
tc = [
    # Auth
    ('TC-01', 'Auth', 'Đăng ký thành công',                  'POST /register dữ liệu hợp lệ', 'Tạo user, redirect /login', 'High'),
    ('TC-02', 'Auth', 'Đăng ký username đã tồn tại',         'username = admin đã có',         'Lỗi "Username đã tồn tại"', 'High'),
    ('TC-03', 'Auth', 'Đăng ký password < 6 ký tự',          'password = "abc"',               'Lỗi validate password',     'Medium'),
    ('TC-04', 'Auth', 'Đăng ký email không hợp lệ',          'email = "abc"',                  'Lỗi "Email không hợp lệ"', 'Medium'),
    ('TC-05', 'Auth', 'Đăng nhập đúng',                      'customer1 / pass123',            'Redirect /, navbar tên user', 'High'),
    ('TC-06', 'Auth', 'Đăng nhập sai password',              'customer1 / wrong',              'Lỗi "Sai thông tin đăng nhập"', 'High'),
    ('TC-07', 'Auth', 'Đăng xuất',                           'Click "Đăng xuất"',              'Session huỷ, redirect /login', 'Medium'),
    # Room
    ('TC-08', 'Room', 'Xem danh sách phòng',                 'GET /rooms',                     'Liệt kê phòng AVAILABLE', 'High'),
    ('TC-09', 'Room', 'Tìm phòng theo ngày',                 'GET /rooms?checkIn=…&checkOut=…','Loại phòng có booking trùng', 'High'),
    ('TC-10', 'Room', 'Lọc theo loại phòng',                 'GET /rooms?type=SUITE',          'Chỉ hiện SUITE', 'Medium'),
    ('TC-11', 'Room', 'Xem chi tiết phòng',                  'GET /rooms/{id} hợp lệ',         'Đầy đủ info + nút Đặt ngay', 'High'),
    # Booking
    ('TC-12', 'Booking', 'Đặt phòng thành công',             'Form hợp lệ + voucher hợp lệ',   'Booking PENDING, total_price đúng', 'High'),
    ('TC-13', 'Booking', 'Đặt phòng trùng lịch',             'Phòng đã có booking trùng',      'Lỗi "Phòng đã được đặt"', 'High'),
    ('TC-14', 'Booking', 'check_out ≤ check_in',             'Ngày sai chiều',                 'Lỗi validate ngày', 'High'),
    ('TC-15', 'Booking', 'check_in trong quá khứ',           'check_in = hôm qua',             'Lỗi "Ngày quá khứ"', 'Medium'),
    ('TC-16', 'Booking', 'Thanh toán MoMo',                  'Chọn MoMo → xem QR → xác nhận', 'payment=PAID, status=CONFIRMED', 'Medium'),
    ('TC-17', 'Booking', 'Huỷ booking',                      'Booking PENDING → Huỷ',          'status=CANCELLED', 'Medium'),
    ('TC-18', 'Booking', 'Huỷ booking đã COMPLETED',         'Booking COMPLETED',              'Lỗi "không thể huỷ"', 'Low'),
    ('TC-19', 'Booking', 'Thanh toán booking của user khác', 'Login customer1, pay booking customer2', 'HTTP 403', 'High'),
    # Auth/Security
    ('TC-20', 'Auth',    'Customer truy cập /admin',         'Login CUSTOMER → GET /admin',    'HTTP 403 Forbidden', 'High'),
    ('TC-21', 'Auth',    '/my-bookings chưa đăng nhập',      'Chưa login → GET /my-bookings',  'Redirect /login', 'Medium'),
    # Admin Room
    ('TC-22', 'Admin', 'Thêm phòng mới',                    'POST /admin/rooms dữ liệu hợp lệ', 'Phòng xuất hiện trong list', 'Medium'),
    ('TC-23', 'Admin', 'Thêm phòng số trùng',               'Số phòng đã tồn tại',              'Lỗi "Số phòng đã tồn tại"', 'Medium'),
    ('TC-24', 'Admin', 'Sửa phòng',                         'Đổi giá → submit',                 'Giá mới được lưu', 'Medium'),
    ('TC-25', 'Admin', 'Xoá phòng có booking',              'Phòng X có booking PENDING',       'Lỗi không cho xoá', 'Medium'),
    ('TC-26', 'Admin', 'Xoá phòng không có booking',        'Phòng Y sạch',                     'Xoá thành công', 'Low'),
    # Admin Booking
    ('TC-27', 'Admin', 'Xem tất cả booking',                'GET /admin/bookings',              'Bảng đầy đủ mọi booking', 'Medium'),
    ('TC-28', 'Admin', 'Lọc booking theo trạng thái',       '?status=PENDING',                  'Chỉ hiện PENDING', 'Low'),
    ('TC-29', 'Admin', 'Cập nhật trạng thái booking',       'PENDING → CONFIRMED',              'Status cập nhật', 'Low'),
    # Performance
    ('TC-30', 'Perf',  'Trang /rooms < 2 giây',             '100 phòng trong DB',               '< 2 giây', 'Low'),
]
add_table(doc, ['ID', 'Module', 'Tên', 'Điều kiện/Bước', 'Kết quả mong đợi', 'Mức độ'],
          tc, [1.3, 1.5, 4, 4, 4, 1.2])

add_heading(doc, '9.3 Quản lý bug', level=2)
add_para(doc,
    'Dự án dùng GitHub Issues làm bug tracker. '
    'Mỗi bug tạo 1 issue với nhãn: bug, severity:critical/major/minor, module:auth/booking/admin. '
    'Template issue gồm: mô tả lỗi, bước tái hiện, kết quả thực tế, kết quả mong đợi, '
    'môi trường (OS, Java, browser), ảnh chụp màn hình.')

add_heading(doc, '9.4 Tiêu chí hoàn thành (Definition of Done)', level=2)
for s in [
    'Tất cả TC High phải Pass.',
    'Không có bug severity:critical còn mở.',
    'Unit test ./mvnw test chạy xanh (0 failures).',
    'Giao diện hiển thị đúng trên Chrome, Firefox ở độ rộng 375px và 1440px.',
    'Admin có thể thực hiện đầy đủ CRUD phòng, booking, voucher, khách hàng.',
]:
    doc.add_paragraph(s, style='List Bullet')

# ===================================================
# LƯU FILE
# ===================================================
import os, sys
out = r'd:/booking/docs/yeu-cau-thiet-ke.docx'
tmp = r'd:/booking/docs/yeu-cau-thiet-ke-new.docx'
doc.save(tmp)
try:
    if os.path.exists(out):
        os.replace(tmp, out)
    else:
        os.rename(tmp, out)
    print('Saved:', out)
except PermissionError:
    print('Saved (tmp):', tmp)
    print('=> Đóng file Word cũ rồi đổi tên file -new.docx thành yeu-cau-thiet-ke.docx')
