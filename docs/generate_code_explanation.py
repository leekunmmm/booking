# -*- coding: utf-8 -*-
"""
Tạo file Word giải thích code theo từng thành viên.
Chạy: python docs/generate_code_explanation.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# ── helpers ─────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x80)

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)

def sub_bullet(text):
    p = doc.add_paragraph(style='List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(10)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def member_banner(name, mssv, role_desc, color_hex='1A5376'):
    """Tạo banner tên thành viên nổi bật."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    pPr.append(shd)
    run = p.add_run(f'  {name}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run2 = p.add_run(f'  —  MSSV: {mssv}  |  {role_desc}')
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)

def file_label(path):
    """Hiển thị đường dẫn file nổi bật."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'File:  {path}')
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0x99)

def tbl(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1A5376')
        tcPr.append(shd)
        for pp in c.paragraphs:
            for r in pp.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, v in enumerate(row):
            r.cells[i].text = v
            for pp in r.cells[i].paragraphs:
                for run in pp.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ════════════════════════════════════════════════════════
# TRANG BÌA
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nTÀI LIỆU GIẢI THÍCH MÃ NGUỒN\n')
run.bold = True; run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Dự án: LuxeHaven Hotel Booking System\n').font.size = Pt(14)

p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run('Nhóm 07  |  Môn: Lập trình & Kiểm thử phần mềm  |  2025–2026\n\n\n').font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# BẢNG PHÂN CÔNG TỔNG HỢP
# ════════════════════════════════════════════════════════
h1('BẢNG PHÂN CÔNG CÔNG VIỆC')
tbl(
    ['Thành viên', 'MSSV', 'Phần phụ trách'],
    [
        ('Bùi Thanh Hậu',       '2251120152', 'Backend: Model, Config, Repository, RoomService, UserService, Database'),
        ('Nguyễn Lê Duy Hoàng', '2251120013', 'Tài liệu: Soạn Word (generate_docx.py), test cases thủ công'),
        ('Duy',  '', 'Frontend: Giao diện Thymeleaf, trang admin (AdminController)'),
        ('Hưng', '', 'Tính năng đặt phòng: BookingService, BookingController, VoucherService'),
        ('Kịp',  '', 'Thiết kế: Sơ đồ UML/DFD/ERD (draw.io), data.sql'),
        ('Thiện','', 'Kiểm thử: AuthService, JUnit tests (AuthServiceTest, BookingServiceTest)'),
    ],
    col_widths=[4.0, 3.0, 10.5]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 1 — BÙI THANH HẬU
# ════════════════════════════════════════════════════════
member_banner('Bùi Thanh Hậu', '2251120152',
              'Backend: Model, Config, Repository, RoomService, UserService, Database',
              '1A5376')

# --- model/User.java ---
h2('1.1  User.java')
file_label('src/main/java/com/luxehaven/booking/model/User.java')
para('Lớp đại diện cho tài khoản người dùng, tương ứng với bảng users trong cơ sở dữ liệu.')
para('Các trường dữ liệu:', bold=True)
bullet('id — Mã định danh duy nhất, tự tăng (1, 2, 3...)')
bullet('username — Tên đăng nhập, không được trùng với người khác')
bullet('password — Mật khẩu đã qua mã hóa BCrypt (không bao giờ lưu mật khẩu gốc)')
bullet('fullName — Họ và tên đầy đủ')
bullet('email — Email, không được trùng trong hệ thống')
bullet('phone — Số điện thoại (có thể để trống)')
bullet('role — Vai trò: CUSTOMER (khách hàng) hoặc ADMIN (quản trị viên)')
bullet('avatarUrl — Đường dẫn ảnh đại diện, ví dụ: /uploads/avatar_5.jpg')
bullet('memberTier — Hạng thành viên: BRONZE → SILVER → GOLD → DIAMOND → VIP')
bullet('createdAt — Ngày giờ tạo tài khoản, tự động gán khi thêm mới')
para('Phương thức:', bold=True)
bullet('isAdmin() — Trả về true nếu role == ADMIN, dùng để kiểm tra quyền nhanh')

# --- model/Room.java ---
h2('1.2  Room.java')
file_label('src/main/java/com/luxehaven/booking/model/Room.java')
para('Lớp đại diện cho một phòng khách sạn, tương ứng với bảng rooms.')
para('Các trường dữ liệu:', bold=True)
bullet('id — Mã phòng, tự tăng')
bullet('roomNumber — Số phòng hiển thị (ví dụ: "101", "302"), không được trùng')
bullet('type — Loại phòng: STANDARD / DELUXE / SUITE')
bullet('pricePerNight — Giá thuê một đêm tính bằng VNĐ')
bullet('capacity — Số người tối đa có thể ở')
bullet('description — Mô tả chi tiết về phòng')
bullet('status — Trạng thái: AVAILABLE (trống), OCCUPIED (đang ở), MAINTENANCE (bảo trì)')
bullet('imagesRaw — Chuỗi đường dẫn ảnh, các ảnh cách nhau bằng dấu phẩy')
sub_bullet('Ví dụ: "/uploads/room101_1.jpg,/uploads/room101_2.jpg"')
bullet('amenitiesRaw — Chuỗi tên tiện nghi cách nhau bằng dấu phẩy')
sub_bullet('Ví dụ: "WIFI,AC,TV,MINIBAR"')
para('Phương thức:', bold=True)
bullet('getImages() — Tách imagesRaw thành danh sách để vòng lặp Thymeleaf hiển thị từng ảnh')
bullet('getAmenities() — Tách amenitiesRaw thành danh sách enum Amenity để hiển thị icon')
bullet('getFirstImage() — Lấy ảnh đầu tiên làm ảnh đại diện (thumbnail) trong danh sách phòng')

# --- model/Booking.java ---
h2('1.3  Booking.java')
file_label('src/main/java/com/luxehaven/booking/model/Booking.java')
para('Lớp đại diện cho một đơn đặt phòng, tương ứng với bảng bookings.')
para('Các trường dữ liệu:', bold=True)
bullet('id — Mã đơn đặt phòng')
bullet('userId — Trỏ đến User (người đặt), là khóa ngoại')
bullet('roomId — Trỏ đến Room (phòng được đặt), là khóa ngoại')
bullet('checkIn / checkOut — Ngày nhận phòng và trả phòng')
bullet('totalPrice — Tổng tiền phải trả (đã trừ voucher nếu có)')
bullet('status — Trạng thái đơn: PENDING → CONFIRMED → COMPLETED / CANCELLED')
bullet('paymentStatus — Trạng thái thanh toán: UNPAID → PENDING_CONFIRMATION → PAID')
bullet('paymentMethod — Cách thanh toán: MOMO hoặc BANK_TRANSFER')
bullet('customerFullName / customerPhone / customerEmail — Thông tin người nhận phòng thực tế')
sub_bullet('Có thể khác với tài khoản đặt (ví dụ: đặt hộ người khác)')
bullet('customerNote — Ghi chú của khách (ví dụ: "Cần giường đôi", "Đến muộn 11 giờ đêm")')
bullet('voucherCode — Mã voucher đã dùng, null nếu không áp dụng')
bullet('room / user — Không lưu vào DB, chỉ gán tạm lúc truy vấn để hiển thị thêm thông tin')
para('Phương thức:', bold=True)
bullet('getNumberOfNights() — Tính số đêm ở = checkOut - checkIn')

# --- model/Voucher.java ---
h2('1.4  Voucher.java')
file_label('src/main/java/com/luxehaven/booking/model/Voucher.java')
para('Lớp đại diện cho mã giảm giá, tương ứng với bảng vouchers.')
para('Các trường dữ liệu:', bold=True)
bullet('id — Mã định danh')
bullet('code — Mã voucher, ví dụ: "SUMMER20". Tự động chuyển thành chữ in hoa')
bullet('discountPercent — Phần trăm giảm giá (từ 1 đến 100)')
bullet('validUntil — Voucher hết hạn sau ngày này')
bullet('maxUses — Số lần dùng tối đa (ví dụ: 100)')
bullet('usedCount — Đã dùng bao nhiêu lần rồi')
bullet('active — Quản trị viên có thể tắt voucher bằng cách đặt false')
bullet('targetTier — Hạng thành viên mới được dùng (null = tất cả mọi người)')
para('Phương thức:', bold=True)
bullet('isUsable() — Trả về true nếu: active=true VÀ chưa hết hạn VÀ usedCount < maxUses')

# --- model Enums ---
h2('1.5  Các Enum')
file_label('src/main/java/com/luxehaven/booking/model/*.java')
para('Enum là kiểu dữ liệu chỉ nhận một trong số các giá trị đã định sẵn, giúp tránh lỗi typo và dễ đọc hơn số nguyên.')
tbl(
    ['File', 'Giá trị', 'Dùng để làm gì'],
    [
        ('Role',          'CUSTOMER, ADMIN',                                             'Xác định quyền hạn người dùng'),
        ('RoomType',      'STANDARD, DELUXE, SUITE',                                    'Phân loại hạng phòng'),
        ('RoomStatus',    'AVAILABLE, OCCUPIED, MAINTENANCE',                           'Trạng thái phòng hiện tại'),
        ('BookingStatus', 'PENDING, CONFIRMED, COMPLETED, CANCELLED',                   'Vòng đời đơn đặt phòng'),
        ('PaymentStatus', 'UNPAID, PENDING_CONFIRMATION, PAID',                         'Tiến trình thanh toán'),
        ('PaymentMethod', 'MOMO, BANK_TRANSFER',                                        'Phương thức thanh toán'),
        ('MemberTier',    'BRONZE, SILVER, GOLD, DIAMOND, VIP',                        'Hạng thành viên theo tổng chi tiêu'),
        ('Amenity',       'WIFI, AC, TV, MINIBAR, JACUZZI, BREAKFAST, BUTLER, POOL, GYM, SPA',
                          'Tiện nghi phòng. Mỗi giá trị có tên hiển thị tiếng Việt và tên icon Bootstrap'),
    ],
    col_widths=[3.5, 6.0, 8.0]
)

# --- Config ---
h2('1.6  Lớp Config')
file_label('src/main/java/com/luxehaven/booking/config/')

h3('SecurityConfig.java')
para('Đăng ký bộ mã hóa mật khẩu BCrypt để dùng chung toàn ứng dụng.')
bullet('BCryptPasswordEncoder(10) — độ mạnh 10 nghĩa là tính 2^10 = 1024 vòng hash')
bullet('Mật khẩu một chiều: không thể giải mã ngược, chỉ so sánh bằng matches()')
bullet('@Bean — Spring tự tạo một instance duy nhất và inject vào AuthService')

h3('AuthInterceptor.java')
para('Chặn và kiểm tra quyền truy cập trước khi xử lý mỗi request.')
bullet('preHandle() chạy trước mọi request, kiểm tra URL cần đăng nhập không')
bullet('URL /admin/**, /bookings/**, /my-bookings → bắt buộc phải đăng nhập')
bullet('URL /admin/** → phải là ADMIN, nếu không trả lỗi 403 Forbidden')
bullet('Chưa đăng nhập → tự động redirect đến /login?redirect=<url gốc>')

h3('WebConfig.java')
para('Đăng ký interceptor và cấu hình phục vụ ảnh tải lên.')
bullet('addInterceptors() — Kết nối AuthInterceptor vào pipeline Spring MVC')
bullet('addResourceHandlers() — Ánh xạ URL /uploads/** → thư mục uploads/ trên ổ đĩa')
sub_bullet('Ờ vậy ảnh phòng upload lên mới truy cập được qua trình duyệt')

h3('DataSeeder.java')
para('Chạy một lần khi ứng dụng khởi động để tạo dữ liệu mặc định.')
bullet('Kiểm tra nếu chưa có admin thì tạo tài khoản admin/admin123')
bullet('Đảm bảo luôn có tài khoản vào quản trị ngay từ lần đầu chạy')

# --- Repository ---
h2('1.7  Lớp Repository')
file_label('src/main/java/com/luxehaven/booking/repository/')
para('Repository là tầng truy cập cơ sở dữ liệu. Dự án dùng JdbcTemplate để viết SQL tường minh thay vì JPA tự động sinh SQL.')

h3('UserRepository.java')
bullet('findByUsername(username) — Tìm user khi đăng nhập')
bullet('findById(id) — Tìm user theo mã')
bullet('findAll() — Lấy toàn bộ danh sách')
bullet('save(user) — Thêm mới nếu id=null, cập nhật nếu id đã có')
bullet('existsByUsername() / existsByEmail() — Kiểm tra trùng khi đăng ký')
bullet('deleteById(id) — Xóa tài khoản')

h3('RoomRepository.java')
bullet('findAll() — Lấy toàn bộ phòng')
bullet('findAllAvailable() — Chỉ lấy phòng đang trống')
bullet('findFiltered(type, status, search) — Tìm kiếm có lọc (dùng trong trang admin)')
bullet('findAvailableInRange(checkIn, checkOut, type) — Tìm phòng không bị đặt trong khoảng ngày này')
bullet('updateStatus(id, status) — Cập nhật trạng thái phòng sau khi booking thay đổi')
bullet('hasActiveBookings(roomId) — Kiểm tra phòng còn booking đang chạy không (để ngăn xóa)')

h3('BookingRepository.java')
bullet('save(booking) — Lưu đơn mới hoặc cập nhật')
bullet('findByUserId(userId) — Lấy tất cả đơn của một khách')
bullet('findByStatus(status) / findAll() — Lấy đơn theo trạng thái hoặc tất cả')
bullet('hasOverlap(...) — Kiểm tra phòng có bị đặt chồng ngày không')
bullet('sumPaidByUserId(userId) — Tổng tiền đã trả của một user (để tính hạng thành viên)')
bullet('countByStatus(status) / totalRevenue() — Số liệu thống kê cho dashboard admin')

h3('VoucherRepository.java')
bullet('findByCode(code) — Tìm voucher theo mã khi khách áp dụng')
bullet('incrementUsedCount(id) — Tăng số lần đã dùng lên 1')
bullet('recordUsage(voucherId, userId) — Ghi lại ai đã dùng voucher này')
bullet('hasUserUsed(voucherId, userId) — Ngăn người dùng dùng cùng voucher hai lần')
bullet('saveVoucher() / removeSavedVoucher() / findSavedByUserId() — Quản lý kho voucher đã lưu của người dùng')

# --- Services ---
h2('1.8  RoomService.java')
file_label('src/main/java/com/luxehaven/booking/service/RoomService.java')
para('Xử lý logic liên quan đến quản lý phòng.')
bullet('search(checkIn, checkOut, type) — Tìm phòng trống:')
sub_bullet('Nếu không nhập ngày → trả về tất cả phòng AVAILABLE/OCCUPIED')
sub_bullet('Nếu có ngày → loại bỏ phòng đang bị đặt chồng lịch')
bullet('create(form) / update(id, form) — Thêm mới / sửa phòng với kiểm tra đầy đủ')
sub_bullet('normalizeImages(): chấp nhận ảnh phân cách bằng dấu phẩy hoặc xuống dòng, chuẩn hóa về dấu phẩy')
bullet('delete(id) — Chỉ xóa được phòng không có booking PENDING hoặc CONFIRMED')

h2('1.9  UserService.java')
file_label('src/main/java/com/luxehaven/booking/service/UserService.java')
para('Quản lý tài khoản và tính hạng thành viên.')
bullet('totalSpent(userId) — Cộng tổng tiền các booking đã PAID của người dùng')
bullet('calculateTier(userId) — Dựa trên tổng chi tiêu để xếp hạng:')
sub_bullet('Dưới 5 triệu → BRONZE  |  5–20 triệu → SILVER  |  20–50 triệu → GOLD')
sub_bullet('50–100 triệu → DIAMOND  |  Trên 100 triệu → VIP')
bullet('countCustomers() — Đếm số tài khoản không phải admin (dùng trong dashboard)')

# --- Schema ---
h2('1.10  schema.sql  &  data.sql')
file_label('src/main/resources/schema.sql  |  data.sql')
para('schema.sql định nghĩa cấu trúc 6 bảng trong SQLite:')
tbl(
    ['Bảng', 'Lưu gì'],
    [
        ('users',              'Tài khoản người dùng'),
        ('rooms',              'Thông tin phòng khách sạn'),
        ('bookings',           'Đơn đặt phòng'),
        ('vouchers',           'Mã giảm giá'),
        ('user_voucher_usage', 'Lịch sử: ai đã dùng voucher nào (ngăn dùng 2 lần)'),
        ('user_saved_vouchers','Kho voucher đã lưu của từng user'),
    ],
    col_widths=[5.0, 12.5]
)
para('')
para('data.sql chèn dữ liệu mẫu: 6 phòng gồm 101, 102 (STANDARD), 201, 202 (DELUXE), 301, 302 (SUITE) với giá từ 500.000 đến 2.200.000 VNĐ/đêm.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 2 — NGUYỄN LÊ DUY HOÀNG
# ════════════════════════════════════════════════════════
member_banner('Nguyễn Lê Duy Hoàng', '2251120013',
              'Tài liệu: Soạn Word, test cases thủ công',
              '1A6B3A')

h2('2.1  generate_docx.py')
file_label('docs/generate_docx.py')
para('Script Python tự động tạo file Word tài liệu yêu cầu và thiết kế (yeu-cau-thiet-ke.docx).')
para('Nội dung tài liệu Word được tạo ra:', bold=True)
bullet('Chương 1 — Mô tả dự án, mục tiêu, đối tượng sử dụng')
bullet('Chương 2 — 20 yêu cầu chức năng (FR) và các yêu cầu phi chức năng')
bullet('Chương 3 — Thiết kế hệ thống: ERD, Class Diagram, DFD, Use Case Diagram')
bullet('Chương 4 — Thiết kế giao diện 18 trang màn hình')
bullet('Chương 5 — Kế hoạch kiểm thử với 30 test cases')
para('Cách chạy:', bold=True)
code('pip install python-docx')
code('python docs/generate_docx.py')
para('Thư viện sử dụng: python-docx — tạo và định dạng file .docx bằng code Python.')

h2('2.2  test-cases.md')
file_label('docs/test-cases.md')
para('File Markdown ghi lại 40 test cases thủ công, chia thành các module:')
tbl(
    ['Module', 'Số TC', 'Nội dung kiểm thử'],
    [
        ('Xác thực',          'TC-01 đến TC-07', 'Đăng ký, đăng nhập, đăng xuất'),
        ('Quản lý phòng',     'TC-08 đến TC-11', 'Xem phòng, tìm kiếm, lọc'),
        ('Đặt phòng',         'TC-12 đến TC-21', 'Đặt phòng, thanh toán, hủy'),
        ('Quản trị (Admin)',  'TC-22 đến TC-29', 'Quản lý phòng, booking'),
        ('Hiệu năng & UI',    'TC-30',            'Thời gian tải trang'),
        ('Hồ sơ cá nhân',    'TC-31 đến TC-40', 'Xem/cập nhật profile, upload ảnh, validate form đăng ký'),
    ],
    col_widths=[4.0, 3.5, 10.0]
)

doc.add_paragraph()
para('Chi tiết test cases hồ sơ cá nhân (do Nguyễn Lê Duy Hoàng thực hiện):', bold=True)
tbl(
    ['ID', 'Tên test', 'Bước thực hiện', 'Kết quả mong đợi', 'Mức độ'],
    [
        ('TC-31', 'Xem trang hồ sơ',
         'Đăng nhập → vào /profile',
         'Hiển thị đúng họ tên, email, phone, hạng thành viên',
         'Medium'),
        ('TC-32', 'Cập nhật họ tên, SĐT',
         'Sửa họ tên + phone hợp lệ → Lưu',
         'Thông tin mới được lưu, hiển thị lại đúng',
         'Medium'),
        ('TC-33', 'Cập nhật email trùng',
         'Đổi email sang email đã có tài khoản khác → Lưu',
         'Lỗi "Email đã được sử dụng"',
         'Medium'),
        ('TC-34', 'Upload ảnh đại diện hợp lệ',
         'Chọn file .jpg < 10MB → Lưu',
         'Ảnh mới hiển thị trên navbar và trang profile',
         'Medium'),
        ('TC-35', 'Upload file không phải ảnh',
         'Chọn file .pdf hoặc .exe → Lưu',
         'Lỗi hoặc không chấp nhận file',
         'Low'),
        ('TC-36', 'Truy cập /profile chưa đăng nhập',
         'GET /profile khi chưa login',
         'Redirect /login?redirect=/profile',
         'High'),
        ('TC-37', 'Đăng ký confirmPassword không khớp',
         'password="abc123", confirmPassword="abc456" → Submit',
         'Lỗi "Mật khẩu xác nhận không khớp"',
         'Medium'),
        ('TC-38', 'Đăng ký họ tên để trống',
         'Để trống trường họ tên → Submit',
         'Lỗi "Họ tên không được để trống"',
         'Medium'),
        ('TC-39', 'Đăng ký username < 3 ký tự',
         'username="ab" → Submit',
         'Lỗi "Tên đăng nhập phải có ít nhất 3 ký tự"',
         'Medium'),
        ('TC-40', 'Đăng nhập lại khi đã login',
         'Đang login → truy cập /login trực tiếp',
         'Redirect về trang chủ',
         'Low'),
    ],
    col_widths=[1.2, 4.0, 5.5, 5.5, 1.8]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 3 — DUY
# ════════════════════════════════════════════════════════
member_banner('Duy', '',
              'Frontend: Giao diện Thymeleaf, trang admin (AdminController)',
              '6B3A8A')

h2('3.1  AdminController.java')
file_label('src/main/java/com/luxehaven/booking/controller/AdminController.java')
para('Controller xử lý toàn bộ trang quản trị tại URL /admin/...')
para('Các chức năng chính:', bold=True)

h3('Dashboard  —  GET /admin')
bullet('Lấy thống kê tổng hợp: tổng số phòng, tổng đơn, doanh thu, số khách hàng')
bullet('Hiển thị trên trang dashboard cho admin thấy toàn cảnh hệ thống')

h3('Quản lý phòng  —  /admin/rooms')
bullet('GET /admin/rooms — Danh sách tất cả phòng, có thể lọc theo loại/trạng thái/tên')
bullet('GET /admin/rooms/new — Hiển thị form thêm phòng mới')
bullet('POST /admin/rooms/create — Nhận dữ liệu form, upload ảnh, lưu phòng vào DB')
sub_bullet('Upload ảnh: đọc file từ MultipartFile, lưu vào thư mục /uploads/, lưu đường dẫn vào imagesRaw')
bullet('GET /admin/rooms/{id}/edit — Form sửa phòng với dữ liệu hiện tại điền sẵn')
bullet('POST /admin/rooms/{id}/update — Cập nhật thông tin phòng')
bullet('POST /admin/rooms/{id}/delete — Xóa phòng (nếu không có booking đang chạy)')

h3('Quản lý đặt phòng  —  /admin/bookings')
bullet('Hiển thị danh sách đơn, lọc theo trạng thái (PENDING/CONFIRMED/...) và trạng thái thanh toán')
bullet('Tìm kiếm theo tên khách hoặc mã đơn')
bullet('POST /admin/bookings/{id}/confirm-payment — Xác nhận đã nhận tiền → PAID + CONFIRMED')
bullet('POST /admin/bookings/{id}/status — Đổi trạng thái đơn thủ công')

h3('Quản lý voucher  —  /admin/vouchers')
bullet('Danh sách voucher, thêm mới, sửa, xóa')
bullet('Mỗi voucher có thể gắn targetTier để chỉ áp dụng cho hạng thành viên nhất định')

h3('Quản lý khách hàng  —  /admin/customers')
bullet('Danh sách tất cả tài khoản CUSTOMER')
bullet('Hiển thị tổng chi tiêu và hạng thành viên hiện tại của từng khách')
bullet('Xem chi tiết lịch sử booking của từng khách')

h2('3.2  HomeController.java')
file_label('src/main/java/com/luxehaven/booking/controller/HomeController.java')
para('Controller cho trang khách hàng (không cần đăng nhập).')
bullet('GET / — Trang chủ: hiển thị phòng nổi bật, thống kê số loại phòng')
bullet('GET /rooms — Danh sách phòng: nhận tham số tìm kiếm từ URL (checkIn, checkOut, type)')
bullet('GET /rooms/{id} — Chi tiết phòng: thông tin đầy đủ + ảnh + tiện nghi + trạng thái sẵn có')
bullet('GET /contact, /services — Trang liên hệ và dịch vụ (trang tĩnh)')

h2('3.3  GlobalModelAttributes.java')
file_label('src/main/java/com/luxehaven/booking/controller/GlobalModelAttributes.java')
para('Class đặc biệt với @ControllerAdvice: tự động chạy cho mọi request, không cần gọi thủ công.')
bullet('Gắn currentUser (người đang đăng nhập) vào model → tất cả template đều dùng được')
bullet('Gắn allAmenities (danh sách 10 tiện nghi) để form phòng hiển thị đầy đủ checkbox')

h2('3.4  Templates Thymeleaf')
file_label('src/main/resources/templates/')
para('Các file HTML dùng Thymeleaf để hiển thị dữ liệu động từ Controller.')
tbl(
    ['File template', 'Trang hiển thị'],
    [
        ('index.html',              'Trang chủ'),
        ('rooms/list.html',         'Danh sách phòng + tìm kiếm'),
        ('rooms/detail.html',       'Chi tiết một phòng'),
        ('auth/login.html',         'Đăng nhập'),
        ('auth/register.html',      'Đăng ký'),
        ('auth/profile.html',       'Trang cá nhân + đổi thông tin'),
        ('booking/new.html',        'Form đặt phòng'),
        ('booking/payment.html',    'Trang hướng dẫn thanh toán QR'),
        ('booking/my-bookings.html','Danh sách đơn của tôi'),
        ('booking/detail.html',     'Chi tiết một đơn đặt phòng'),
        ('admin/dashboard.html',    'Dashboard admin'),
        ('admin/rooms/*.html',      'Quản lý phòng (danh sách, form thêm/sửa)'),
        ('admin/bookings/*.html',   'Quản lý đặt phòng'),
        ('admin/vouchers/*.html',   'Quản lý voucher'),
        ('admin/customers/*.html',  'Quản lý khách hàng'),
    ],
    col_widths=[5.5, 12.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 4 — HƯNG
# ════════════════════════════════════════════════════════
member_banner('Hưng', '',
              'Tính năng đặt phòng: BookingService, BookingController, VoucherService',
              '8A3A1A')

h2('4.1  BookingService.java')
file_label('src/main/java/com/luxehaven/booking/service/BookingService.java')
para('Service xử lý toàn bộ vòng đời của một đơn đặt phòng.')

h3('createBooking(userId, form)  —  @Transactional')
para('@Transactional nghĩa là: nếu có lỗi xảy ra ở bất kỳ bước nào, toàn bộ thao tác sẽ bị hủy (rollback), không lưu dữ liệu nửa vời.')
bullet('Bước 1 — Kiểm tra dữ liệu đầu vào:')
sub_bullet('checkOut phải sau checkIn, không đặt ngày trong quá khứ')
sub_bullet('Phải chọn phương thức thanh toán hợp lệ (MOMO hoặc BANK_TRANSFER)')
sub_bullet('Họ tên, số điện thoại, email người nhận phòng không được để trống')
bullet('Bước 2 — Kiểm tra phòng còn trống không (hasOverlap):')
sub_bullet('Truy vấn DB xem phòng đó có booking PENDING/CONFIRMED chồng ngày không')
sub_bullet('Nếu có → báo lỗi "Phòng đã được đặt trong khoảng thời gian này"')
bullet('Bước 3 — Tính tiền: số đêm × giá/đêm')
bullet('Bước 4 — Nếu có voucher: gọi VoucherService.validateAndGet() → trừ % giảm giá')
bullet('Bước 5 — Lưu booking vào DB: status=PENDING, paymentStatus=UNPAID')

h3('pay(bookingId, userId)')
bullet('Khách nhấn "Tôi đã thanh toán" sau khi chuyển khoản')
bullet('Kiểm tra: đúng booking của user này, chưa hủy, chưa gửi xác nhận trước đó')
bullet('Đổi paymentStatus → PENDING_CONFIRMATION (chờ admin duyệt)')

h3('adminConfirmPayment(bookingId)')
bullet('Admin xác nhận đã nhận tiền thực tế')
bullet('Đổi paymentStatus → PAID, status → CONFIRMED')

h3('cancel(bookingId, userId)')
bullet('Không hủy được nếu: đã hoàn tất (COMPLETED), đã đến ngày nhận phòng')
bullet('Sau khi hủy: kiểm tra phòng còn booking khác không → nếu không thì chuyển phòng về AVAILABLE')

h3('adminUpdateStatus(bookingId, newStatus)')
bullet('Admin đổi trạng thái thủ công')
bullet('Tự động đồng bộ paymentStatus và trạng thái phòng theo trạng thái mới')

h2('4.2  BookingController.java')
file_label('src/main/java/com/luxehaven/booking/controller/BookingController.java')
para('Controller nhận request từ browser và gọi BookingService.')
bullet('GET /bookings/new?roomId=X — Hiển thị form đặt phòng cho phòng X')
sub_bullet('Kiểm tra phòng tồn tại, lấy thông tin phòng để điền sẵn vào form')
bullet('POST /bookings/create — Nhận dữ liệu form, gọi BookingService.createBooking()')
sub_bullet('Nếu thành công → chuyển đến trang thanh toán')
sub_bullet('Nếu lỗi → hiển thị lại form với thông báo lỗi')
bullet('GET /my-bookings — Lấy danh sách đơn của người dùng đang đăng nhập')
bullet('GET /bookings/{id} — Xem chi tiết một đơn')
bullet('POST /bookings/{id}/pay — Khách xác nhận đã chuyển tiền')
bullet('POST /bookings/{id}/cancel — Khách hủy đơn')
bullet('GET /bookings/{id}/payment — Trang hướng dẫn quét QR (MoMo hoặc ngân hàng)')

h2('4.3  VoucherService.java')
file_label('src/main/java/com/luxehaven/booking/service/VoucherService.java')
para('Service quản lý toàn bộ logic liên quan đến mã giảm giá.')

h3('validateAndGet(code, userId)  —  Kiểm tra voucher trước khi dùng')
bullet('Tìm voucher theo code trong DB, nếu không có → lỗi "Mã không hợp lệ"')
bullet('Gọi isUsable(): active=true, chưa hết hạn, usedCount < maxUses')
bullet('Nếu voucher có targetTier: tính hạng thành viên hiện tại của user → phải khớp')
sub_bullet('Ví dụ: voucher GOLD chỉ dành cho thành viên GOLD. User SILVER dùng sẽ bị từ chối')
bullet('Kiểm tra user này chưa từng dùng voucher đó (hasUserUsed)')

h3('use(voucher, userId)  —  Sau khi đặt phòng thành công')
bullet('Tăng usedCount lên 1')
bullet('Ghi lại vào bảng user_voucher_usage: người này đã dùng voucher này rồi')

h3('Kho voucher cá nhân')
bullet('saveForUser() — User lưu voucher vào kho để dùng sau')
bullet('getSavedByUser() — Lấy danh sách voucher đã lưu của user')
bullet('removeSavedForUser() — Xóa khỏi kho sau khi dùng')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 5 — KỊP
# ════════════════════════════════════════════════════════
member_banner('Kịp', '',
              'Thiết kế: Sơ đồ UML/DFD/ERD (draw.io), dữ liệu mẫu',
              '5A5A1A')

h2('5.1  Sơ đồ Use Case  —  usecase.drawio')
file_label('docs/diagrams/usecase.drawio')
para('Mô tả các chức năng của hệ thống và ai có thể thực hiện chúng.')
bullet('3 tác nhân: Guest (khách vãng lai), Customer (đã đăng nhập), Admin')
bullet('Customer kế thừa quyền của Guest (generalization)')
bullet('19 use case từ UC-01 đến UC-19:')
sub_bullet('Guest: xem phòng, tìm kiếm, xem chi tiết, xem giá')
sub_bullet('Customer: đăng ký, đăng nhập, đặt phòng, thanh toán, hủy đơn, xem lịch sử, quản lý voucher, hồ sơ')
sub_bullet('Admin: quản lý phòng, quản lý đơn, quản lý voucher, quản lý khách hàng, xác nhận thanh toán')

h2('5.2  Sơ đồ ERD  —  erd.drawio')
file_label('docs/diagrams/erd.drawio')
para('Sơ đồ quan hệ thực thể, cho thấy cấu trúc CSDL và mối quan hệ giữa các bảng.')
bullet('users 1 → nhiều bookings (một người có thể đặt nhiều phòng)')
bullet('rooms 1 → nhiều bookings (một phòng có thể được đặt nhiều lần)')
bullet('bookings tham chiếu vouchers qua cột voucher_code (không phải khóa ngoại cứng)')
bullet('4 bảng chính: users, rooms, bookings, vouchers')
bullet('2 bảng phụ: user_voucher_usage, user_saved_vouchers')

h2('5.3  Sơ đồ DFD  —  dfd.drawio')
file_label('docs/diagrams/dfd.drawio')
para('Sơ đồ luồng dữ liệu, cho thấy dữ liệu chạy qua hệ thống như thế nào.')
bullet('Mức 0 (Context): hệ thống là một hộp đen, nhận input từ Khách/Admin, trả output ra')
bullet('Mức 1: chi tiết 5 luồng xử lý chính:')
sub_bullet('P1 — Quản lý tài khoản (đăng ký, đăng nhập, cập nhật hồ sơ)')
sub_bullet('P2 — Tìm kiếm và xem phòng')
sub_bullet('P3 — Đặt phòng và thanh toán')
sub_bullet('P4 — Khách hàng quản lý booking cá nhân')
sub_bullet('P5 — Quản trị viên quản lý hệ thống')

h2('5.4  Sơ đồ Class Diagram  —  classdiagram.drawio')
file_label('docs/diagrams/classdiagram.drawio')
para('Sơ đồ lớp UML, cho thấy cấu trúc code theo hướng đối tượng.')
bullet('4 lớp chính: User, Room, Booking, Voucher — tương ứng 4 bảng trong DB')
bullet('8 enum: Role, RoomType, RoomStatus, BookingStatus, PaymentStatus, PaymentMethod, MemberTier, Amenity')
bullet('Mũi tên dependency từ Booking → Room, Booking → User, Booking → Voucher')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHẦN 6 — THIỆN
# ════════════════════════════════════════════════════════
member_banner('Thiện', '',
              'Kiểm thử: AuthService, JUnit tests (AuthServiceTest, BookingServiceTest)',
              '1A5A5A')

h2('6.1  AuthService.java')
file_label('src/main/java/com/luxehaven/booking/service/AuthService.java')
para('Service xử lý đăng ký và xác thực đăng nhập.')

h3('register(form)')
para('Xử lý đăng ký tài khoản mới. Kiểm tra từng điều kiện, ném lỗi ngay nếu không đạt:')
bullet('Username ≥ 3 ký tự, nếu không → lỗi "Tên đăng nhập phải có ít nhất 3 ký tự"')
bullet('Password ≥ 6 ký tự, nếu không → lỗi "Mật khẩu phải có ít nhất 6 ký tự"')
bullet('confirmPassword phải trùng với password, nếu không → lỗi "Mật khẩu xác nhận không khớp"')
bullet('Email phải có dấu @, nếu không → lỗi "Email không hợp lệ"')
bullet('fullName không được trống')
bullet('Kiểm tra username chưa có trong DB (existsByUsername), nếu trùng → lỗi')
bullet('Kiểm tra email chưa có trong DB (existsByEmail), nếu trùng → lỗi')
bullet('Nếu qua hết: mã hóa password bằng BCrypt rồi lưu vào DB')

h3('authenticate(username, password)')
para('Xử lý đăng nhập:')
bullet('Tìm user theo username trong DB')
bullet('Nếu không tìm thấy → trả về Optional.empty() (đăng nhập thất bại)')
bullet('Dùng passwordEncoder.matches(password_nhập, password_hash_trong_DB) để so sánh')
bullet('Nếu sai password → trả về Optional.empty()')
bullet('Nếu đúng → trả về Optional.of(user) (đăng nhập thành công)')
para('Lý do dùng Optional thay vì null: buộc Controller phải kiểm tra trước khi dùng, tránh NullPointerException.')

h2('6.2  AuthServiceTest.java')
file_label('src/test/java/com/luxehaven/booking/service/AuthServiceTest.java')
para('File kiểm thử tự động cho AuthService, dùng JUnit 5 và Mockito.')
para('Mockito là gì: thay thế các dependency thật (UserRepository, BCrypt) bằng bản giả để test Service độc lập, không cần DB thật.')
para('Các test case:', bold=True)
bullet('register_success — Khi dữ liệu hợp lệ: phải gọi userRepository.save() 1 lần, trả về User')
bullet('register_usernameExists — Khi username đã có: existsByUsername trả true → phải ném BusinessException')
bullet('register_passwordTooShort — Password 4 ký tự → phải ném BusinessException ngay, không gọi save()')
bullet('register_emailInvalid — Email không có @ → phải ném BusinessException')
bullet('authenticate_success — Username/password đúng: BCrypt.matches trả true → trả Optional có user')
bullet('authenticate_wrongPassword — Password sai: BCrypt.matches trả false → trả Optional.empty()')

h2('6.3  BookingServiceTest.java')
file_label('src/test/java/com/luxehaven/booking/service/BookingServiceTest.java')
para('File kiểm thử tự động cho BookingService, dùng JUnit 5 và Mockito.')
para('Các test case:', bold=True)
bullet('createBooking_success — Form hợp lệ, phòng không bị chồng lịch → Booking trả về có status PENDING')
bullet('createBooking_checkoutBeforeCheckin — checkOut ≤ checkIn → phải ném BusinessException')
bullet('createBooking_roomNotFound — roomId không tồn tại trong DB → phải ném BusinessException')
bullet('createBooking_overlap — Phòng bị chồng lịch (hasOverlap trả true) → phải ném BusinessException')
bullet('cancel_success — Hủy đơn PENDING → status phải chuyển thành CANCELLED, gọi bookingRepository.save()')
bullet('cancel_alreadyCancelled — Hủy đơn đã CANCELLED rồi → phải ném BusinessException')

h2('6.4  BookingApplicationTests.java')
file_label('src/test/java/com/luxehaven/booking/BookingApplicationTests.java')
para('Test đơn giản nhất: kiểm tra Spring khởi động được không.')
bullet('contextLoads() — Nếu test này pass, nghĩa là toàn bộ cấu hình Spring Boot không có lỗi')
bullet('Đây là bước kiểm tra cơ bản nhất trước khi chạy các test khác')

h2('6.5  Cách chạy kiểm thử')
para('Chạy toàn bộ test tự động:')
code('mvn test')
para('Xem kết quả chi tiết:')
code('target/surefire-reports/TEST-*.xml')
para('Ngoài test tự động, còn có 30 test case thủ công trong docs/test-cases.md — cần chạy ứng dụng thật và kiểm tra từng bước.')

doc.add_page_break()

# ════════════════════════════════════════════════════════
# PHỤ LỤC — LUỒNG HOẠT ĐỘNG
# ════════════════════════════════════════════════════════
h1('PHỤ LỤC — LUỒNG ĐẶT PHÒNG TỔNG HỢP')
para('Cho thấy các phần của từng thành viên phối hợp với nhau như thế nào:')
doc.add_paragraph()
tbl(
    ['Bước', 'Ai thực hiện', 'Code xử lý'],
    [
        ('1. Khách tìm phòng trống',           'Duy (HomeController)',      'HomeController.rooms() → RoomService.search()'),
        ('2. RoomService tìm phòng',            'Hậu (RoomService)',         'RoomRepository.findAvailableInRange()'),
        ('3. Khách điền form đặt phòng',        'Duy (template)',            'booking/new.html'),
        ('4. Submit form đặt phòng',            'Hưng (BookingController)',  'BookingController.createBooking()'),
        ('5. Kiểm tra và lưu đơn',              'Hưng (BookingService)',     'BookingService.createBooking() @Transactional'),
        ('6. Kiểm tra voucher (nếu có)',        'Hưng (VoucherService)',     'VoucherService.validateAndGet()'),
        ('7. Khách quét QR thanh toán',         'Duy (template)',            'booking/payment.html'),
        ('8. Khách gửi xác nhận đã chuyển',    'Hưng (BookingService)',     'BookingService.pay() → PENDING_CONFIRMATION'),
        ('9. Admin xem danh sách chờ duyệt',   'Duy (AdminController)',     'AdminController.bookings()'),
        ('10. Admin xác nhận thanh toán',       'Hưng (BookingService)',     'BookingService.adminConfirmPayment() → PAID'),
    ],
    col_widths=[5.5, 4.5, 7.5]
)

# ── Lưu file ────────────────────────────────────────────
output = os.path.join(os.path.dirname(__file__), 'giai-thich-code-v2.docx')
doc.save(output)
print(f'Da tao: {output}')
